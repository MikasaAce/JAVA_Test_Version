import os

from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm, Inches
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import datetime
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from xml.sax.saxutils import escape
import traceback
from matplotlib import font_manager
import io

from app.api.config import config as data_class

# 获取py 文件所在目录
current_path = os.path.dirname(__file__)

# 把这个目录设置成工作目录
os.chdir(current_path)

csv_path = data_class.pdf_save_path


class Graphs_word:
    @staticmethod
    def draw_title(doc, title: str, font_size: int, style=None, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   bold=True, line_spacing=None):
        """
        通用的标题/文本绘制函数

        Args:
            doc: Word文档对象
            title: 文本内容
            font_size: 字体大小
            alignment: 对齐方式，默认为居中
            bold: 是否加粗，默认为True
            style: 样式名称，默认为None（自动根据字体大小选择）
            line_spacing: 行间距，默认为None（自动根据字体大小设置）
        """
        # 设置默认样式
        if style is None:
            if font_size >= 24:
                style = 'Heading 1' if font_size >= 32 else 'Heading 2'
            else:
                style = 'Normal'

        # 设置默认行间距
        if line_spacing is None:
            if font_size >= 32:
                line_spacing = Pt(50)
            elif font_size >= 20:
                line_spacing = Pt(30)
            else:
                line_spacing = Pt(30)

        # 添加段落并设置基本属性
        p = doc.add_paragraph()
        p.alignment = alignment
        p.style = doc.styles[style]

        # 添加文字
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # 统一格式设置
        run.font.bold = bold
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = False

        # 设置段落格式
        # p.paragraph_format.line_spacing = line_spacing
        return p

    # 保留原有的便捷方法作为快捷方式
    @staticmethod
    def draw_title_32(doc, title: str):
        return Graphs_word.draw_title(doc, title, 32)

    @staticmethod
    def draw_text_32(doc, title: str):
        return Graphs_word.draw_title(doc, title, 32, 'Normal')

    @staticmethod
    def draw_title_24(doc, title: str):
        return Graphs_word.draw_title(doc, title, 24)

    @staticmethod
    def draw_text_24(doc, title: str):
        return Graphs_word.draw_title(doc, title, 24,'Normal')

    @staticmethod
    def draw_title_20(doc, title: str):
        return Graphs_word.draw_title(doc, title, 20)

    @staticmethod
    def draw_text_14(doc, title: str):
        return Graphs_word.draw_title(doc, title, 14)

    @staticmethod
    def draw_left_16(doc, title: str):
        return Graphs_word.draw_title(doc, title, 16, 'Heading 3', WD_ALIGN_PARAGRAPH.LEFT, True)

    @staticmethod
    def draw_left_14(doc, title: str):
        return Graphs_word.draw_title(doc, title, 14, 'Heading 4', WD_ALIGN_PARAGRAPH.LEFT, True)

    @staticmethod
    def draw_text_12(doc, text: str):
        p = Graphs_word.draw_title(doc, text, 12, 'Normal', WD_ALIGN_PARAGRAPH.LEFT, False)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        return p

    @staticmethod
    def draw_text_12_bold(doc, text: str):
        p = Graphs_word.draw_title(doc, text, 12, 'Normal', WD_ALIGN_PARAGRAPH.LEFT, True)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        return p

    @staticmethod
    def draw_text_12_overview(doc, text: str):
        p = Graphs_word.draw_title(doc, text, 12, 'Normal', WD_ALIGN_PARAGRAPH.LEFT, False)
        p.paragraph_format.space_before = Pt(10)
        return p


    def set_cell_border(cell, color_hex):
        """设置单元格边框颜色"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 创建边框元素
        borders = OxmlElement('w:tcBorders')

        # 定义边框类型
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), color_hex)
            borders.append(border)

        # 将边框元素添加到单元格属性中
        tcPr.append(borders)

    def set_cell_style(cell, font_color, alignment):
        """设置单元格样式"""
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = font_color

                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    def draw_table(doc, data):
        # 创建表格
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.autofit = False

        # 设置行高
        for row in table.rows:
            row.height = Mm(8)

        # 计算可用宽度
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # 定义列宽比例并设置列宽
        width_ratios = [1.1, 2.5, 1.1, 1.7]
        total_ratio = sum(width_ratios)
        for i, ratio in enumerate(width_ratios):
            table.columns[i].width = int(available_width * ratio / total_ratio)

        # 合并单元格
        if len(data) == 10:
            table.cell(8, 1).merge(table.cell(8, 3))
            table.cell(9, 1).merge(table.cell(9, 3))

        # 填充数据
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)

                # 设置单元格格式
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                        # 对偶数列应用加粗和底纹
                        if j % 2 == 0:
                            run.font.bold = True
                            tc_pr = cell._tc.get_or_add_tcPr()
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="e6e6e6"/>')
                            tc_pr.append(shading)

                # 设置垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders") or OxmlElement('w:tblBorders')
        if tblBorders is not None and tblBorders not in tblPr:
            tblPr.append(tblBorders)

        # 设置各方向边框
        for border_name in ['top', 'bottom', 'left', 'right']:
            border = tblBorders.find(qn(f'w:{border_name}')) or OxmlElement(f'w:{border_name}')
            if border not in tblBorders:
                tblBorders.append(border)

            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '16')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')

        # # 创建表格
        # table = doc.add_table(rows=len(data), cols=2)
        #
        # # 设置列宽
        # table.columns[0].width = Mm(24)
        # table.columns[1].width = Mm(60)
        #
        # # 设置行高
        # for row in table.rows:
        #     row.height = Mm(13)
        #
        # # 填充数据并设置样式
        # for i, (col1, col2) in enumerate(data):
        #     row_cells = table.rows[i].cells
        #     cell0, cell1 = row_cells[0], row_cells[1]
        #
        #     # 设置单元格内容
        #     cell0.text, cell1.text = col1, col2
        #
        #     # 处理第一列样式
        #     for paragraph in cell0.paragraphs:
        #         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #         paragraph.paragraph_format.space_before = Mm(0)
        #         paragraph.paragraph_format.space_after = Mm(0)
        #         for run in paragraph.runs:
        #             run.font.name = 'SimSun'
        #             run.font.color.rgb = RGBColor(47, 79, 79)
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        #
        #     # 设置第一列背景颜色
        #     shading = OxmlElement('w:shd')
        #     shading.set(qn('w:fill'), 'd5dae6')
        #     cell0._tc.get_or_add_tcPr().append(shading)
        #
        #     # 处理第二列样式
        #     for paragraph in cell1.paragraphs:
        #         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #         paragraph.paragraph_format.space_before = Mm(0)
        #         paragraph.paragraph_format.space_after = Mm(0)
        #         for run in paragraph.runs:
        #             run.font.name = 'Times New Roman'
        #             run.font.color.rgb = RGBColor(47, 79, 79)
        #
        #     # 设置垂直居中
        #     cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #
        # # 设置所有单元格边框
        # for row in table.rows:
        #     for cell in row.cells:
        #         Graphs_word.set_cell_border(cell, '808080')

    @staticmethod
    def draw_table_vul_level(doc, data):
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        table.autofit = False

        table.rows[0].height = Mm(8)
        table.rows[1].height = Mm(15)
        table.rows[2].height = Mm(15)
        table.rows[3].height = Mm(15)

        # 计算可用宽度
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # 定义列宽比例并设置列宽
        width_ratios = [0.9, 1, 5.2]
        total_ratio = sum(width_ratios)
        for i, ratio in enumerate(width_ratios):
            table.columns[i].width = int(available_width * ratio / total_ratio)

        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)

                # 设置单元格格式
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    for run in para.runs:

                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                        # 为第一行设置加粗和底纹
                        if i == 0:
                            run.font.bold = True
                            tc_pr = cell._tc.get_or_add_tcPr()
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="e6e6e6"/>')
                            tc_pr.append(shading)
                    if i < 1 or j < 2:
                        para.alignment = WD_ALIGN_VERTICAL.CENTER  # 设置水平居中
                    if i < 1 or j >= 2:
                        para.paragraph_format.space_before = Pt(5)

                # 设置垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders") or OxmlElement('w:tblBorders')
        if tblBorders is not None and tblBorders not in tblPr:
            tblPr.append(tblBorders)

        # 设置各方向边框
        for border_name in ['top', 'bottom', 'left', 'right']:
            border = tblBorders.find(qn(f'w:{border_name}')) or OxmlElement(f'w:{border_name}')
            if border not in tblBorders:
                tblBorders.append(border)

            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '16')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')

    def draw_table_1(doc, data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.autofit = False

        # 设置行高
        for row in table.rows:
            row.height = Mm(8)

        # 计算可用宽度
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # 定义列宽比例并设置列宽
        width_ratios = [1.3, 2.5, 0.6, 0.6, 0.6, 0.7]
        total_ratio = sum(width_ratios)
        for i, ratio in enumerate(width_ratios):
            table.columns[i].width = int(available_width * ratio / total_ratio)

        # 填充数据
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)

                shading_colors = {
                    0: "e6e6e6",  # 前两列使用相同颜色
                    1: "e6e6e6",  # 前两列使用相同颜色
                    2: "ff0000",  # 红色
                    3: "ffc000",  # 橙色
                    4: "0070c0",  # 蓝色
                    5: "e6e6e6",  # 灰色
                }

                # 设置单元格格式
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                        # 为第一行设置加粗和底纹
                        if i == 0:
                            run.font.bold = True

                            # 只在第一次运行时添加底纹，避免重复添加
                            if j in shading_colors and not hasattr(cell, '_shading_added'):
                                tc_pr = cell._tc.get_or_add_tcPr()
                                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_colors[j]}"/>')
                                tc_pr.append(shading)
                                # 标记已添加底纹，避免重复操作
                                cell._shading_added = True

                # 设置垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # 设置水平居中
                if j >= 2:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_VERTICAL.CENTER

        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders") or OxmlElement('w:tblBorders')
        if tblBorders is not None and tblBorders not in tblPr:
            tblPr.append(tblBorders)

            # 设置各方向边框
        for border_name in ['top', 'bottom', 'left', 'right']:
            border = tblBorders.find(qn(f'w:{border_name}')) or OxmlElement(f'w:{border_name}')
            if border not in tblBorders:
                tblBorders.append(border)

            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '16')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')

    def draw_table_1_before(doc, data):
        # 创建表格
        table = doc.add_table(rows=len(data), cols=3)

        # 设置列宽
        table.columns[0].width = Mm(15)
        table.columns[1].width = Mm(60)
        table.columns[2].width = Mm(15)

        # 设置行高
        for row in table.rows:
            row.height = Mm(13)

        # 填充数据并设置样式
        for i, (col1, col2, col3) in enumerate(data):
            row_cells = table.rows[i].cells
            cell0, cell1, cell2 = row_cells[0], row_cells[1], row_cells[2]

            # 设置单元格内容
            cell0.text, cell1.text, cell2.text = col1, col2, col3

            # 设置第一行背景颜色
            if i == 0:
                for cell in row_cells:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'd5dae6')
                    cell._tc.get_or_add_tcPr().append(shading)

            # 设置单元格样式
            for cell in row_cells:
                Graphs_word.set_cell_style(cell, RGBColor(47, 79, 79), WD_ALIGN_PARAGRAPH.CENTER)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 设置所有单元格边框
        for row in table.rows:
            for cell in row.cells:
                Graphs_word.set_cell_border(cell, '808080')

    def draw_table_2(doc, data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.autofit = False

        # 设置行高
        for row in table.rows:
            row.height = Mm(8)

        # 计算可用宽度
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        # 定义列宽比例并设置列宽
        width_ratios = [3.4, 0.8, 0.8, 0.8, 0.8]
        total_ratio = sum(width_ratios)
        for i, ratio in enumerate(width_ratios):
            table.columns[i].width = int(available_width * ratio / total_ratio)

        # 填充数据
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)

                shading_colors = {
                    0: "99ccff",  # 浅蓝色
                    1: "ff0000",  # 红色
                    2: "ffc000",  # 橙色
                    3: "0070c0",  # 蓝色
                    4: "e6e6e6",  # 灰色
                }

                # 设置单元格格式
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                        # 为第一行设置加粗和底纹
                        if i == 0:
                            run.font.bold = True

                            # 只在第一次运行时添加底纹，避免重复添加
                            if j in shading_colors and not hasattr(cell, '_shading_added'):
                                tc_pr = cell._tc.get_or_add_tcPr()
                                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_colors[j]}"/>')
                                tc_pr.append(shading)
                                # 标记已添加底纹，避免重复操作
                                cell._shading_added = True

                # 设置垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # 设置水平居中
                if j >= 1:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_VERTICAL.CENTER

        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders") or OxmlElement('w:tblBorders')
        if tblBorders is not None and tblBorders not in tblPr:
            tblPr.append(tblBorders)

            # 设置各方向边框
        for border_name in ['top', 'bottom', 'left', 'right']:
            border = tblBorders.find(qn(f'w:{border_name}')) or OxmlElement(f'w:{border_name}')
            if border not in tblBorders:
                tblBorders.append(border)

            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '16')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')

    def draw_table_3(doc, data, id, location):
        # 过滤字符
        if '爆发行代码片段' in data:
            # 分割为多行
            lines = data['爆发行代码片段'].split('\n')

            # 清理每行：去除首尾空白，并移除 "line X" 模式的内容
            cleaned_lines = []
            for line in lines:
                cleaned_line = line.strip()
                # 使用正则表达式移除 "line 数字" 模式
                import re
                cleaned_line = re.sub(r'>>>', '', cleaned_line)
                cleaned_line = re.sub(r'Line\s+\d+:', '', cleaned_line)
                cleaned_lines.append(cleaned_line)

            # 重新组合为字符串
            data['爆发行代码片段'] = '\n'.join(cleaned_lines)

        # 创建表格并设置基本样式
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        table.autofit = False

        # 设置行高
        table.rows[0].height = Mm(12)
        table.rows[1].height = Mm(8)

        # 计算可用宽度并设置列宽
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin

        width_ratios = [4.2, 1, 1]
        total_ratio = sum(width_ratios)
        for i, ratio in enumerate(width_ratios):
            table.columns[i].width = int(available_width * ratio / total_ratio)

        # 合并单元格
        table.cell(0, 0).merge(table.cell(0, 2))
        table.cell(2, 0).merge(table.cell(2, 2))

        # 准备单元格内容和样式配置
        cells_config = [
            {
                'cell': table.cell(0, 0),
                'text': f'序号{id}：{data["缺陷路径"]}',
                'font_size': Pt(10),
                'bold': True,
                'shading': '99ccff',
                'vertical_align': True
            },
            {
                'cell': table.cell(1, 0),
                'text': f'缺陷名称：{data["缺陷名称"]}',
                'font_size': Pt(10),
                'bold': True,
                'vertical_align': True
            },
            {
                'cell': table.cell(1, 1),
                'text': '缺陷等级',
                'font_size': Pt(10),
                'bold': True,
                'alignment': WD_ALIGN_VERTICAL.CENTER,
                'shading': 'e2efd9',
                'vertical_align': True
            },
            {
                'cell': table.cell(1, 2),
                'text': data['缺陷等级'],
                'font_size': Pt(10),
                'bold': True,
                'alignment': WD_ALIGN_VERTICAL.CENTER,
                'vertical_align': True
            }
        ]

        # 添加行号
        if not location:
            location = 0
        lines = data['爆发行代码片段'].split('\n')
        # 方法1：使用模糊查找
        burst_point = data['爆发点'].strip()
        for idx, line in enumerate(lines):
            if burst_point in line.strip():
                highlight_idx = idx
                break
        else:
            # 如果没找到，使用原始方法并捕获异常
            highlight_idx = -1  # 或 lines.index(burst_point)
        start_idx = location - highlight_idx
        numbered_lines = [f'{start_idx + i}  {line}' for i, line in enumerate(lines)]
        numbered_code = '\n'.join(numbered_lines)

        parts1 = numbered_code.split(numbered_lines[highlight_idx], 1)
        strings1 = [
            '爆发行代码片段：\n',
            parts1[0],
            f'{location}' + data['爆发点'],
            parts1[1],
        ]

        cells_config.append({
            'cell': table.cell(2, 0),
            'text_parts': strings1,
            'font_size': Pt(10),
            'bold_parts': [True, False, False, False],
            'color_parts': [None, None, RGBColor(255, 0, 16), None],
            'vertical_align': True
        })

        # 应用样式到所有单元格
        for config in cells_config:
            cell = config['cell']

            # 设置垂直居中
            if config.get('vertical_align'):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

            # 设置单元格背景色
            if 'shading' in config:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{config["shading"]}"/>')
                tc_pr.append(shading)

            # 处理普通文本单元格
            if 'text' in config:
                cell.text = config['text']

                for para in cell.paragraphs:
                    # 设置段落对齐
                    if 'alignment' in config:
                        para.alignment = config['alignment']

                    for run in para.runs:
                        run.font.size = config['font_size']
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                        run.font.bold = config.get('bold', False)

            # 处理多部分文本单元格
            elif 'text_parts' in config:
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

                for i, text in enumerate(config['text_parts']):
                    run = para.add_run(text)
                    run.font.size = config['font_size']
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    run.font.bold = config['bold_parts'][i]

                    if 'color_parts' in config and config['color_parts'][i]:
                        run.font.color.rgb = config['color_parts'][i]

    def draw_bar(doc, data):
        # 设置中英文字体
        plt.rcParams.update({
            'font.sans-serif': ['SimSun', 'SimHei', 'Microsoft YaHei'],
            'font.family': 'sans-serif',
            'axes.unicode_minus': False
        })

        # 创建图形
        plt.figure(figsize=(12, 9))

        categories, repaired_data, unrepaired_data, false_positive_data = data
        colors = ['#a4cbb9', '#f7685b', '#9E9E9E']  # 绿色:已修复, 橙色:未修复, 灰色:误报

        # 创建堆叠柱状图
        bars1 = plt.bar(categories, repaired_data, color=colors[0], width=0.4,
                        label='已修复', edgecolor='white', linewidth=0.5)
        bars2 = plt.bar(categories, unrepaired_data, bottom=repaired_data, width=0.4,
                        color=colors[1], label='未修复', edgecolor='white', linewidth=0.5)
        bottom_values = [sum(x) for x in zip(repaired_data, unrepaired_data)]
        bars3 = plt.bar(categories, false_positive_data, bottom=bottom_values, color=colors[2], width=0.4,
                        label='误报', edgecolor='white', linewidth=0.5)

        # 添加数据标签
        for i, (bar1, bar2, bar3) in enumerate(zip(bars1, bars2, bars3)):
            # 第一个系列标签
            if (height := bar1.get_height()) > 0:
                plt.text(bar1.get_x() + bar1.get_width() / 2, height / 2, f'{int(height)}',
                         ha='center', va='center', fontsize=17, fontweight='bold', color='white')

            # 第二个系列标签
            if (height := bar2.get_height()) > 0:
                y_pos = repaired_data[i] + height / 2
                plt.text(bar2.get_x() + bar2.get_width() / 2, y_pos, f'{int(height)}',
                         ha='center', va='center', fontsize=17, fontweight='bold', color='white')

            # 第三个系列标签
            if (height := bar3.get_height()) > 0:
                y_pos = bottom_values[i] + height / 2
                plt.text(bar3.get_x() + bar3.get_width() / 2, y_pos, f'{int(height)}',
                         ha='center', va='center', fontsize=17, fontweight='bold', color='white')

        # 设置图表样式
        plt.legend(loc='best', frameon=True, fancybox=True, shadow=True, fontsize=15)
        plt.grid(axis='y', alpha=0.3, linestyle='--')
        plt.xticks(rotation=0, ha='center')
        plt.yticks(fontsize=17)

        # 设置坐标轴样式
        ax = plt.gca()
        for spine in ['top', 'right', 'left']:
            ax.spines[spine].set_visible(False)

        ax.tick_params(axis='x', which='major', direction='out', length=6,
                       width=1.2, color='gray', labelsize=17)

        # 设置Y轴范围
        max_value = max(sum(x) for x in zip(repaired_data, unrepaired_data, false_positive_data))
        plt.ylim(0, max_value * 1.18)
        plt.tight_layout()

        # 保存并插入文档
        image_stream = io.BytesIO()
        plt.savefig(image_stream, format='png', dpi=300, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
        plt.close()
        image_stream.seek(0)

        doc.add_picture(image_stream, width=Inches(6))
        image_stream.close()

    # 综合评分算法
    def calculate_comprehensive_score(high_risk, medium_risk, low_risk, total_lines):
        """
        综合评分算法
        :param high_risk: 高危缺陷数
        :param medium_risk: 中危缺陷数
        :param low_risk: 低危缺陷数
        :param total_lines: 代码总行数
        :param file_count: 文件数量
        :return: 综合评分(0-100)
        """

        # 1. 缺陷密度得分（30%权重）
        defect_density = (high_risk + medium_risk + low_risk) / total_lines * 1000
        density_score = max(0, 100 - defect_density * 20)

        # 2. 风险严重性得分（50%权重）
        risk_weighted = high_risk * 3 + medium_risk * 2 + low_risk * 1
        max_acceptable_risk = total_lines / 1000  # 每千行代码可接受风险基数
        risk_score = max(0, 100 - (risk_weighted / max(max_acceptable_risk, 1)) * 30)

        # 3. 缺陷分布得分（20%权重）
        # 检查是否多种类型缺陷并存
        risk_variety = len([x for x in [high_risk, medium_risk, low_risk] if x > 0])
        distribution_score = max(60, 100 - (risk_variety - 1) * 10) if risk_variety > 0 else 100

        # 综合计算
        final_score = (
                density_score * 0.3 +
                risk_score * 0.5 +
                distribution_score * 0.2
        )

        return round(max(0, min(100, final_score)), 1)

    # 风险评级算法
    def risk_rating(score, high_risk_count):
        """
        风险评级算法
        :param score: 综合评分
        :param high_risk_count: 高危缺陷数量
        :return: 风险等级
        """
        if score >= 90 and high_risk_count == 0:
            return "优秀"
        elif score >= 80 and high_risk_count <= 5:
            return "良好"
        elif score >= 70:
            return "中等"
        elif score >= 50:
            return "一般"
        elif score >= 30:
            return "中危"
        else:
            return "高危"

    def export_word(itemName, word_Time,  vuls, file_location, final_results, number, version, risk_level_dict, bar_info, pie_info, git_info, overview_info):
        try:
            doc = Document()
            # 添加标题,第一页
            Graphs_word.draw_title_32(doc, '静态代码安全扫描报告')
            Graphs_word.draw_text_24(doc, '\n')
            Graphs_word.draw_text_32(doc, '\n\n')
            Graphs_word.draw_title_20(doc, '项目名称：{}'.format(itemName))
            Graphs_word.draw_text_32(doc, '\n\n')
            Graphs_word.draw_title_20(doc, '报告生成时间：{}'.format(word_Time))
            doc.add_page_break()

            # 第二页，生成项目表格
            Graphs_word.draw_title_24(doc, '第一部分：漏洞扫描结果汇总')
            Graphs_word.draw_left_16(doc, '1.项目基本情况')
            # 添加表格
            if git_info[0]:
                overview_info.extend([
                    ('git仓库地址', f'{git_info[0] or ""}'),
                    ('分支', f'{git_info[1] or "master"}'),
                ])
            Graphs_word.draw_table(doc, overview_info)
            high_vul = int(pie_info[1][0])
            med_vul = int(pie_info[1][1])
            low_vul = int(pie_info[1][2])
            code_lines = int(overview_info[3][3][:-1])
            score = Graphs_word.calculate_comprehensive_score(high_vul,med_vul,low_vul,code_lines)
            risk_rating = Graphs_word.risk_rating(score, high_vul)
            text=f'''本次对于 {itemName} 源代码工程的扫描测试，扫描了 {overview_info[2][1]} 个文件，其中 {code_lines} 行代码。本次检测语言为：{overview_info[1][1]}。共发现了 {overview_info[6][1]} 个缺陷，其中高危 {high_vul} 个，中危 {med_vul} 个，低危 {low_vul} 个。检测结果综合评分为 {score}分，综合风险评级为 {risk_rating}。风险分布如下所示:'''

            Graphs_word.draw_text_12_overview(doc, f'''    {text}''')

            risk_data = [
                ['缺陷等级', '缺陷数量', '缺陷等级评定标准'],
                ['高危', f'{pie_info[1][0]}', '''1) 可直接导致严重的逻辑漏洞（如任意账户密码重置等），攻击者能够直接获取服务器权限或敏感数据；
2) 可直接导致严重的信息泄露，如数据库连接字符串、密钥文件、大量用户敏感信息明文传输等。'''],
                ['中危', f'{pie_info[1][1]}', '''1) 可能间接导致安全风险，通常需要特定条件或用户交互才能利用；
2) 可能导致一般性的信息泄露，如服务器版本信息、非关键路径下的文件路径泄露等。'''],
                ['低危', f'{pie_info[1][2]}', '''1) 属于代码质量或安全规范问题，短期内几乎无法被直接利用于攻击；
2) 非常轻微的信息泄露，如普通的客户端错误信息回显。'''],
            ]

            Graphs_word.draw_table_vul_level(doc, risk_data)

            # 生成漏洞表格
            Graphs_word.draw_left_16(doc, '2.漏洞检测结果汇总')
            Graphs_word.draw_left_14(doc, '2.1 检测结果汇总表')
            Graphs_word.draw_text_14(doc, '漏洞检测结果详情统计')
            data = [('序号', '缺陷名称', '高危', '中危', '低危', '总计')]
            for index, (vul_name, description, repairPlan, file_list) in enumerate(final_results, start=1):
                high_vul = 0
                medium_vul = 0
                low_vul = 0
                for (fileName,vulType,location,source_code,repair_code,Sink,Enclosing_Method,Source, filepath,risk_level) in file_list:
                    if risk_level=='高危':
                        high_vul += 1
                    elif risk_level=='中危':
                        medium_vul += 1
                    elif risk_level=='低危':
                        low_vul += 1
                data.append([index, vul_name, high_vul, medium_vul, low_vul, high_vul+medium_vul+low_vul])
            Graphs_word.draw_table_1(doc, data)

            # 生成图表
            Graphs_word.draw_left_14(doc, '2.2 漏洞类型统计图')
            # 生成漏洞统计柱状图（要体现修复情况，已修复、未修复、误报等）
            Graphs_word.draw_text_14(doc, '漏洞统计柱状图')
            Graphs_word.draw_bar(doc, bar_info)
            doc.add_page_break()


            # 图表之后的详细记录
            Graphs_word.draw_title_24(doc, '第二部分：漏洞扫描结果详情')
            if version is not None:
                if version == 'Developer Workbook':
                    j = 1
                    vuls = vuls.split(',')
                    for i, vul in enumerate(vuls):
                        vul = vul.strip()
                        level = risk_level_dict[vul]
                        Description = None
                        suggestion = None
                        Graphs_word.draw_left_16(doc, '{id} {cweName}'.format(id=j, cweName=vul))
                        Graphs_word.draw_left_14(doc, '{id}.1 漏洞概述'.format(id=j))
                        # 添加漏洞的信息表格
                        data = [
                            ('漏洞编号', '漏洞名称', '风险等级'),
                            ('{}'.format(i + 1), '{}'.format(vul), '{}'.format(level))
                        ]
                        Graphs_word.draw_table_1_before(doc, data)
                        Graphs_word.draw_left_14(doc, '{id}.2 漏洞详情'.format(id=j))
                        num = 0
                        t = 1
                        for i2 in range(0, len(number)):
                            if vul.lower().replace(" ", "") == number[i2][0].replace(" ", ""):
                                for i3 in range(num, num + int(number[i2][1])):
                                    Graphs_word.draw_text_12_bold(doc, '{id}）漏洞定位'.format(id=t))
                                    Graphs_word.draw_text_12(
                                        doc,
                                        '{fileName}:{location}'.format(
                                            fileName=file_location[i3][0], location=file_location[i3][2]
                                        )
                                    )
                                    if file_location[i3][2]:
                                        Graphs_word.draw_text_12(doc, '漏洞分析')
                                        Graphs_word.draw_text_12(doc, '漏洞爆发点:{Sink}'.format(Sink=file_location[i3][5]))
                                        Graphs_word.draw_text_12(
                                            doc,
                                            '爆发点函数:{Enclosing_Method}'.format(Enclosing_Method=file_location[i3][6]))
                                        Graphs_word.draw_text_12(doc, '缺陷源:{Source}'.format(Source=file_location[i3][7]))
                                    if file_location[i3][3]:
                                        Graphs_word.draw_text_12(doc, '漏洞源代码')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][3]))
                                    if file_location[i3][4]:
                                        Graphs_word.draw_text_12(doc, '修复参考')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][4]))
                                    t += 1
                            num += int(number[i2][1])
                        if Description:
                            Graphs_word.draw_left_14(doc, '{id}.3 漏洞描述'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(Description))
                        if suggestion:
                            Graphs_word.draw_left_14(doc, '{id}.4 修复建议'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(suggestion))
                        j += 1
                elif 'OWASP' in version:
                    if '2010' in version:
                        table = pd.read_csv(os.path.join(csv_path, "csv/OWASP_2010.csv"), encoding='gbk')
                    elif '2013' in version:
                        table = pd.read_csv(os.path.join(csv_path, "csv/OWASP_2013.csv"), encoding='gbk')
                    elif '2017' in version:
                        table = pd.read_csv(os.path.join(csv_path, "csv/OWASP_2017.csv"), encoding='gbk')
                    vul_names = table['漏洞名称'].values.tolist()
                    levels = table['漏洞等级'].values.tolist()
                    Descriptions = table['漏洞描述'].values.tolist()
                    suggestions = table['修复方案'].values.tolist()
                    j = 1
                    vuls = vuls.split(',')
                    for vul in vuls:
                        id = 0
                        for i, _ in enumerate(vul_names):
                            if _ == vul:
                                id = i
                        vul_name = vul_names[id]
                        level = levels[id]
                        Description = Descriptions[id]
                        suggestion = suggestions[id]
                        Graphs_word.draw_left_16(doc, '{id} {cweName}'.format(id=j, cweName=vul))
                        Graphs_word.draw_left_14(doc, '{id}.1 漏洞概述'.format(id=j))
                        # 添加漏洞的信息表格
                        data = [
                            ('漏洞编号', '漏洞名称', '风险等级'),
                            ('{}'.format(i + 1), '{}'.format(vul), '{}'.format(level))
                        ]
                        Graphs_word.draw_table_1_before(doc, data)
                        Graphs_word.draw_left_14(doc, '{id}.2 漏洞详情'.format(id=j))
                        num = 0
                        t = 1
                        for i2 in range(0, len(number)):
                            if vul.lower().replace(" ", "") == number[i2][0].replace(" ", ""):
                                for i3 in range(num, num + int(number[i2][1])):
                                    Graphs_word.draw_text_12_bold(doc, '{id}）漏洞定位'.format(id=t))
                                    Graphs_word.draw_text_12(
                                        doc,
                                        '{fileName}:{location}'.format(
                                            fileName=file_location[i3][0], location=file_location[i3][2]
                                        )
                                    )
                                    if file_location[i3][2]:
                                        Graphs_word.draw_text_12(doc, '漏洞分析')
                                        Graphs_word.draw_text_12(doc, '漏洞爆发点:{Sink}'.format(Sink=file_location[i3][5]))
                                        Graphs_word.draw_text_12(
                                            doc,
                                            '爆发点函数:{Enclosing_Method}'.format(Enclosing_Method=file_location[i3][6]))
                                        Graphs_word.draw_text_12(doc, '缺陷源:{Source}'.format(Source=file_location[i3][7]))
                                    if file_location[i3][3]:
                                        Graphs_word.draw_text_12(doc, '漏洞源代码')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][3]))
                                    if file_location[i3][4]:
                                        Graphs_word.draw_text_12(doc, '修复参考')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][4]))
                                    t += 1
                            num += int(number[i2][1])
                        if Description:
                            Graphs_word.draw_left_14(doc, '{id}.3 漏洞描述'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(Description))
                        if suggestion:
                            Graphs_word.draw_left_14(doc, '{id}.4 修复建议'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(suggestion))
                        j += 1
                elif 'CWE' in version:
                    table = pd.read_csv(os.path.join(csv_path, "csv/漏洞知识库.csv"), encoding='gbk')
                    ids = table['漏洞ID'].values.tolist()
                    cweNames = table['漏洞名称'].values.tolist()
                    levels = table['漏洞等级'].values.tolist()
                    Descriptions = table['漏洞描述'].values.tolist()
                    suggestions = table['修复方案'].values.tolist()
                    j = 1
                    vuls = vuls.split(',')
                    for vul in vuls:
                        id = 0
                        for i, _ in enumerate(ids):
                            if _ == vul.lower().strip():
                                id = i
                        cweName = cweNames[id]
                        level = levels[id]
                        Description = Descriptions[id]
                        suggestion = suggestions[id]
                        Graphs_word.draw_left_16(doc, '{id} {cweName}'.format(id=j, cweName=vul))
                        Graphs_word.draw_left_14(doc, '{id}.1 漏洞概述'.format(id=j))
                        # 添加漏洞的信息表格
                        data = [
                            ('漏洞编号', '漏洞名称', '风险等级'),
                            ('{}'.format(i + 1), '{}'.format(vul), '{}'.format(level))
                        ]
                        Graphs_word.draw_table_1_before(doc, data)
                        Graphs_word.draw_left_14(doc, '{id}.2 漏洞详情'.format(id=j))
                        num = 0
                        t = 1
                        for i2 in range(0, len(number)):
                            if vul.lower().replace(" ", "") == number[i2][0].replace(" ", ""):
                                for i3 in range(num, num + int(number[i2][1])):
                                    Graphs_word.draw_text_12_bold(doc, '{id}）漏洞定位'.format(id=t))
                                    Graphs_word.draw_text_12(
                                        doc,
                                        '{fileName}:{location}'.format(
                                            fileName=file_location[i3][0], location=file_location[i3][2]
                                        )
                                    )
                                    if file_location[i3][2]:
                                        Graphs_word.draw_text_12(doc, '漏洞分析')
                                        Graphs_word.draw_text_12(doc, '漏洞爆发点:{Sink}'.format(Sink=file_location[i3][5]))
                                        Graphs_word.draw_text_12(
                                            doc,
                                            '爆发点函数:{Enclosing_Method}'.format(Enclosing_Method=file_location[i3][6]))
                                        Graphs_word.draw_text_12(doc, '缺陷源:{Source}'.format(Source=file_location[i3][7]))
                                    if file_location[i3][3]:
                                        Graphs_word.draw_text_12(doc, '漏洞源代码')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][3]))
                                    if file_location[i3][4]:
                                        Graphs_word.draw_text_12(doc, '修复参考')
                                        Graphs_word.draw_text_12(doc, '{location}'.format(location=file_location[i3][4]))
                                    t += 1
                            num += int(number[i2][1])
                        if Description:
                            Graphs_word.draw_left_14(doc, '{id}.3 漏洞描述'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(Description))
                        if suggestion:
                            Graphs_word.draw_left_14(doc, '{id}.4 修复建议'.format(id=j))
                            Graphs_word.draw_text_12(doc, '{}'.format(suggestion))
                        j += 1
            else:
                for index, (vul_name, description, repairPlan, file_list) in enumerate(final_results, start=1):
                    description = '\n'.join([line.strip().replace("\u200b", "") for line in description.split('\n') if line.strip()])
                    repairPlan = '\n'.join([line.strip().replace("\u200b", "") for line in repairPlan.split('\n') if line.strip()])

                    Graphs_word.draw_left_16(doc, '{id} {vul_name}'.format(id=index, vul_name=vul_name))
                    Graphs_word.draw_left_14(doc, '{id}.1 缺陷描述'.format(id=index))
                    Graphs_word.draw_text_12(doc, '{}'.format(description))
                    Graphs_word.draw_left_14(doc, '{id}.2 解决方案'.format(id=index))
                    Graphs_word.draw_text_12(doc, '{}'.format(repairPlan))
                    Graphs_word.draw_left_14(doc, '{id}.3 缺陷统计'.format(id=index))
                    high_vul = 0
                    medium_vul = 0
                    low_vul = 0
                    for (
                    fileName, vulType, location, source_code, repair_code, Sink, Enclosing_Method, Source, filepath,
                    risk_level) in file_list:
                        if risk_level == '高危':
                            high_vul += 1
                        elif risk_level == '中危':
                            medium_vul += 1
                        elif risk_level == '低危':
                            low_vul += 1
                    table2_data = [('缺陷名称', '高危', '中危', '低危', '总计')]
                    table2_data.append((vul_name,high_vul, medium_vul, low_vul, high_vul+medium_vul+low_vul))
                    Graphs_word.draw_table_2(doc, table2_data)
                    Graphs_word.draw_left_14(doc, '{id}.4 代码位置'.format(id=index))
                    file_id = 0
                    for fileName,vulType,location,source_code,repair_code,Sink,Enclosing_Method,Source, filepath,risk_level in file_list:
                        if risk_level:
                            file_id+=1
                            table3_data = {
                                '缺陷路径': filepath,
                                '缺陷名称': vulType,
                                '缺陷等级': risk_level,
                                '爆发点': Sink,
                                '爆发行代码片段': source_code
                            }
                            Graphs_word.draw_table_3(doc, table3_data, file_id, int(location))
                            doc.add_paragraph()
                        else:
                            continue

            # ================== 保存文件 ==================
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{itemName}_检测报告_{timestamp}.docx"

            doc.save(os.path.join(csv_path, filename))
            print(f"报告已生成：{filename} in {csv_path}/{filename}.xlsx")
            return filename
        except Exception as e:
            traceback.print_exc()
            print(f'导出word文件时出错: {e}')

